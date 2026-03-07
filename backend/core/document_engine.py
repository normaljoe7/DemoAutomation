from docx import Document
from docxtpl import DocxTemplate
import os
import re
import zipfile
import tempfile
import shutil
import subprocess
import platform


def _convert_docx_to_pdf_libreoffice(docx_path: str, output_dir: str):
    """
    Convert a DOCX file to PDF using LibreOffice headless mode.
    Returns the PDF path on success, or None if LibreOffice is not found / conversion fails.
    """
    candidates = []
    if platform.system() == "Windows":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice 7\program\soffice.exe",
            r"C:\Program Files\LibreOffice 6\program\soffice.exe",
        ]
    else:
        import shutil as _shutil
        for name in ("libreoffice", "soffice"):
            found = _shutil.which(name)
            if found:
                candidates.append(found)
                break

    soffice = None
    for c in candidates:
        if os.path.isabs(c):
            if os.path.exists(c):
                soffice = c
                break
        else:
            soffice = c
            break

    if not soffice:
        return None

    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", docx_path, "--outdir", output_dir],
            timeout=120,
            capture_output=True,
            text=True,
        )
        if result.returncode == 0:
            base = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(output_dir, base + ".pdf")
            if os.path.exists(pdf_path):
                return pdf_path
    except Exception:
        pass

    return None

class DocumentEngine:
    def __init__(self):
        pass

    def scan_variables(self, template_path):
        """Scans a DOCX or PPTX file for placeholders in {{variable}} format."""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
            
        if template_path.lower().endswith('.pptx'):
            return self._scan_variables_pptx(template_path)
            
        doc = Document(template_path)

        variables = []
        
        def add_matches(text):
            matches = re.findall(r'\{\{\s*(\w+)\s*\}\}', text)
            for match in matches:
                if match not in variables:
                    variables.append(match)
        
        for paragraph in doc.paragraphs:
            add_matches(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        add_matches(paragraph.text)
                        
        return variables

    def _scan_variables_pptx(self, template_path):
        variables = []
        with zipfile.ZipFile(template_path, 'r') as zin:
            for item in zin.infolist():
                if item.filename.endswith('.xml'):
                    content = zin.read(item.filename).decode('utf-8', errors='ignore')
                    # Find {{var}} even if split across XML tags
                    content = re.sub(r'(\{\{)(?:<[^>]+>)*\s*(\w+)\s*(?:<[^>]+>)*(\}\})', r'\1\2\3', content)
                    matches = re.findall(r'\{\{\s*(\w+)\s*\}\}', content)
                    for match in matches:
                        if match not in variables:
                            variables.append(match)
        return variables

    def fix_template_xml(self, template_path):
        """
        Attempts to fix DOCX XML where placeholders are split across multiple runs.
        e.g. <w:t>{{</w:t>...<w:t>var</w:t>...<w:t>}}</w:t> -> <w:t>{{var}}</w:t>
        """
        temp_dir = tempfile.mkdtemp()
        temp_zip = os.path.join(temp_dir, "temp.docx")
        
        try:
            # Check if we need fixing by reading XML
            with zipfile.ZipFile(template_path, 'r') as zin:
                xml = zin.read('word/document.xml').decode('utf-8')
                
            # Regex to find split tags: {{...}}
            # Pattern: {{ followed by any number of tags/text, ending with }}
            # We want to consolidate.
            # Simple heuristic: Remove XML tags between {{ and }}
            # Careful not to remove formatting tags if possible... 
            # But usually variables don't need mixed formatting.
            
            # Find {{...}} where there are tags in between
            new_xml = re.sub(r'(\{\{)(?:<[^>]+>)*\s*(\w+)\s*(?:<[^>]+>)*(\}\})', r'\1\2\3', xml)
            
            if new_xml != xml:
                # Rewrite the zip
                with zipfile.ZipFile(template_path, 'r') as zin, zipfile.ZipFile(temp_zip, 'w') as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename == 'word/document.xml':
                            zout.writestr(item, new_xml)
                        else:
                            zout.writestr(item, data)
                            
                # Replace original
                shutil.move(temp_zip, template_path)
                return True
        except Exception as e:
            print(f"Failed to fix XML: {e}")
            
        shutil.rmtree(temp_dir, ignore_errors=True)
        return False

    def generate(self, template_path, data, output_path, convert_pdf=False):
        """Generates a document by replacing placeholders with data."""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        if template_path.lower().endswith('.pptx'):
            return self._generate_pptx(template_path, data, output_path, convert_pdf)
            
        try:
            doc = DocxTemplate(template_path)
            # Create a context where keys are matched strictly
            # Note: docxtpl expects keys to match EXACTLY what's inside {{ }}.
            # Our scan_variables extracts 'key' from '{{ key }}'.
            # data has keys that match perfectly.
            
            doc.render(data)
            doc.save(output_path)
            
            pdf_path = None
            error_msg = None

            if convert_pdf:
                abs_output_path = os.path.abspath(output_path)
                output_dir = os.path.dirname(abs_output_path)
                pdf_target = abs_output_path.replace(".docx", ".pdf")

                # Strategy 1 (Windows): docx2pdf via MS Word COM automation
                # Most reliable on business Windows machines with Word installed
                if platform.system() == "Windows":
                    import sys, io
                    original_stdout = sys.stdout
                    original_stderr = sys.stderr
                    try:
                        try:
                            import pythoncom
                            pythoncom.CoInitialize()
                        except Exception:
                            pass
                        from docx2pdf import convert as _docx2pdf_convert
                        dummy_stream = io.StringIO()
                        sys.stdout = dummy_stream
                        sys.stderr = dummy_stream
                        _docx2pdf_convert(abs_output_path, pdf_target)
                        if os.path.exists(pdf_target):
                            pdf_path = pdf_target
                    except Exception as e:
                        print(f"Word/docx2pdf conversion failed: {e}")
                    finally:
                        sys.stdout = original_stdout
                        sys.stderr = original_stderr

                # Strategy 2: LibreOffice headless (cross-platform fallback)
                if not pdf_path:
                    lo_pdf = _convert_docx_to_pdf_libreoffice(abs_output_path, output_dir)
                    if lo_pdf and os.path.exists(lo_pdf):
                        pdf_path = lo_pdf

                if not pdf_path:
                    error_msg = "PDF conversion failed: requires MS Word or LibreOffice"
                    print(error_msg)
                    
            return output_path, pdf_path, error_msg

        except Exception as e:
            raise RuntimeError(f"Document generation failed: {e}")

    def _generate_pptx(self, template_path, data, output_path, convert_pdf=False):
        temp_dir = tempfile.mkdtemp()
        temp_zip = os.path.join(temp_dir, "temp.pptx")
        
        pdf_path = None
        error_msg = None
        
        try:
            with zipfile.ZipFile(template_path, 'r') as zin, zipfile.ZipFile(temp_zip, 'w') as zout:
                for item in zin.infolist():
                    content = zin.read(item.filename)
                    if item.filename.endswith('.xml'):
                        xml_str = content.decode('utf-8', errors='ignore')
                        
                        # Consolidate split tags: 
                        xml_str = re.sub(r'(\{\{)(?:<[^>]+>)*\s*(\w+)\s*(?:<[^>]+>)*(\}\})', r'\1\2\3', xml_str)
                        
                        # Replace
                        for k, v in data.items():
                            placeholder = f"{{{{{k}}}}}"
                            # XML escape
                            val_str = str(v).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            xml_str = xml_str.replace(placeholder, val_str)
                            
                        zout.writestr(item, xml_str.encode('utf-8'))
                    else:
                        zout.writestr(item, content)
            
            shutil.copy(temp_zip, output_path)
            
            if convert_pdf:
                # PDF conversion for PPTX using comtypes (requires pywin32 / MS Office)
                try:
                    import comtypes.client
                    import pythoncom
                    pythoncom.CoInitialize()
                    
                    powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
                    powerpoint.Visible = 1
                    
                    abs_output_path = os.path.abspath(output_path)
                    pdf_target = abs_output_path.replace(".pptx", ".pdf")
                    
                    deck = powerpoint.Presentations.Open(abs_output_path)
                    deck.SaveAs(pdf_target, 32) # 32 is ppSaveAsPDF
                    deck.Close()
                    powerpoint.Quit()
                    
                    pdf_path = pdf_target
                except Exception as e:
                    error_msg = f"PPTX to PDF Conversion Failed: {str(e)}"
                    print(error_msg)
            
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
            
        return output_path, pdf_path, error_msg
